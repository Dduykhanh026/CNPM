
from docx import Document
from collections import Counter

path = 'KH33_SRS_-v3.0.docx'
doc = Document(path)

styles = Counter(p.style.name if p.style is not None else 'None' for p in doc.paragraphs)
headings = [p.text for p in doc.paragraphs if p.style is not None and p.style.name.startswith('Heading') and p.text.strip()]

lines = []
lines.append('Top styles:')
for style, count in styles.most_common(10):
    lines.append(f'{style}: {count}')

lines.append('')
lines.append('Headings:')
for heading in headings:
    lines.append(heading)

lines.append('')
lines.append('First 120 paragraphs:')
for p in doc.paragraphs[:120]:
    style_name = p.style.name if p.style is not None else 'None'
    lines.append(f'[{style_name}] {p.text}')

with open('inspect_output.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(lines))

print('Wrote inspect_output.txt')
